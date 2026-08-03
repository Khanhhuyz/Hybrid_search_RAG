"""
Document Processor Service
Handles text extraction from PDF, DOCX, TXT, and Markdown files.
Integrates TextCleaner for preprocessing before chunking.
"""
import re
import logging
from pathlib import Path
from typing import Optional, List

from app.services.text_cleaner import TextCleaner

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Extract, clean, and normalize text from multiple document formats."""

    def __init__(self):
        self.cleaner = TextCleaner()

    def extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract and clean text from a document based on its type."""
        extractors = {
            ".pdf":  self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt":  self._extract_txt,
            ".md":   self._extract_markdown,
        }
        extractor = extractors.get(file_type.lower())
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Extract raw text (and pages for PDFs)
        if file_type.lower() == ".pdf":
            raw_text, pages = self._extract_pdf_with_pages(file_path)
            # Clean with page-aware header/footer detection
            cleaned = self.cleaner.clean(raw_text, pages=pages)
        else:
            raw_text = extractor(file_path)
            cleaned = self.cleaner.clean(raw_text)

        return cleaned

    # ─── Format-Specific Extractors ──────────────────────────────────────────

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF (without per-page info)."""
        text, _ = self._extract_pdf_with_pages(file_path)
        return text

    def _extract_pdf_with_pages(self, file_path: Path) -> tuple:
        """Extract text from PDF, returning (full_text, list_of_page_texts)."""
        try:
            import pdfplumber
            pages_text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
            return "\n\n".join(pages_text), pages_text
        except ImportError:
            raise RuntimeError("pdfplumber is required for PDF processing. Run: pip install pdfplumber")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

    def _extract_docx(self, file_path: Path) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)
        except ImportError:
            raise RuntimeError("python-docx is required for DOCX processing. Run: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_txt(self, file_path: Path) -> str:
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        for enc in encodings:
            try:
                return file_path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Unable to decode text file: {file_path}")

    def _extract_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown, preserving structure."""
        raw = self._extract_txt(file_path)
        # Strip markdown syntax but keep readable text
        text = re.sub(r"```[\s\S]*?```", "", raw)           # remove code blocks
        text = re.sub(r"`[^`]+`", "", text)                  # remove inline code
        text = re.sub(r"!\[.*?\]\(.*?\)", "", text)          # remove images
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text) # links → text
        text = re.sub(r"#{1,6}\s+", "", text)                # headings
        text = re.sub(r"(\*{1,2}|_{1,2})(.*?)\1", r"\2", text) # bold/italic
        return text.strip()

    # ─── Text Normalization (legacy, now handled by TextCleaner) ─────────────

    def normalize(self, text: str) -> str:
        """Clean and normalize extracted text (legacy compatibility)."""
        return self.cleaner.clean(text)
